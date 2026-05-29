from __future__ import annotations

from collections import defaultdict
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.models import Candidate, Criterion, Template
from app.scoring import summarize_candidate


BLUE = RGBColor(37, 70, 74)
TEAL = RGBColor(22, 105, 122)
LOGO_PATH = Path(__file__).resolve().parent / "assets" / "logo-sie.png"


def percent(value: float) -> str:
    return f"{round((value or 0) * 100, 2)}%"


def score_text(score: float | None) -> str:
    return "Sin evaluar" if score is None else f"{score:.1f}/5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(0)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Title", "Heading 1", "Heading 2"):
        styles[name].font.name = "Aptos Display"
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Superintendencia de Electricidad - Reporte de evaluación curricular")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


def heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = BLUE if level == 1 else TEAL
    return paragraph


def body(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    return paragraph


def bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    return paragraph


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True)
        set_cell_shading(table.rows[0].cells[index], "DCECEA")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    document.add_paragraph()


def add_cover(document: Document, candidate: Candidate, template: Template, summary: dict) -> None:
    if LOGO_PATH.exists():
        logo = document.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.add_run().add_picture(str(LOGO_PATH), width=Inches(2.4))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Reporte individual de evaluación curricular")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(21)
    run.font.color.rgb = BLUE

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(candidate.name)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL

    rows = [
        ["Plantilla de evaluación", template.name],
        ["Cédula / ID", candidate.document_id or "No registrado"],
        ["Evaluador asignado", candidate.evaluator or "No registrado"],
        ["Fecha de generación", datetime.now().strftime("%d/%m/%Y %H:%M")],
        ["Resultado global", percent(summary["global_score"])],
        ["Recomendación", summary["recommendation"]],
    ]
    add_table(document, ["Dato", "Detalle"], rows)

    body(
        document,
        "Este documento resume la evaluación registrada en la plataforma VALCV. Su propósito es apoyar la revisión humana "
        "del expediente curricular y facilitar una lectura ordenada de los criterios, evidencias y resultados obtenidos.",
    )
    document.add_page_break()


def category_narrative(document: Document, template: Template, summary: dict) -> None:
    heading(document, "Resumen ejecutivo de resultados")
    body(
        document,
        f"El candidato fue evaluado con la plantilla \"{template.name}\" y obtuvo un resultado global de "
        f"{percent(summary['global_score'])}. La recomendación calculada por la plataforma es: "
        f"{summary['recommendation']}. Este resultado se deriva de los criterios ponderados definidos en la plantilla y "
        "de las puntuaciones registradas durante la revisión del expediente.",
    )
    if template.description:
        body(document, f"Descripción de la plantilla: {template.description}")

    rows = [[category.name, percent(category.weight), percent(summary["categories"].get(category.name, 0))] for category in template.categories]
    add_table(document, ["Categoría", "Peso definido", "Resultado del candidato"], rows)


def weights_narrative(document: Document, template: Template, criteria: list[Criterion]) -> None:
    heading(document, "Estructura de evaluación")
    body(
        document,
        "La evaluación se organizó en categorías con distintos niveles de importancia dentro del resultado final. "
        "A continuación se presenta la estructura aplicada para valorar el perfil del candidato.",
    )
    grouped: dict[str, list[Criterion]] = defaultdict(list)
    for criterion in criteria:
        grouped[criterion.category].append(criterion)
    for category in template.categories:
        heading(document, f"{category.name} ({percent(category.weight)})", 2)
        children = grouped.get(category.name, [])
        if not children:
            body(document, "No se registraron criterios en esta categoría.")
            continue
        for criterion in children:
            if criterion.is_critical:
                body(document, f"{criterion.aspect}. Requisito de cumplimiento obligatorio.")
            else:
                body(document, f"{criterion.aspect}. Peso relativo dentro de la categoría: {percent(criterion.within_category_weight)}.")


def evaluation_narrative(document: Document, candidate: Candidate, criteria: list[Criterion]) -> None:
    heading(document, "Análisis por criterio")
    body(
        document,
        "Esta sección resume las puntuaciones y comentarios registrados para cada criterio evaluado. Cuando existan "
        "documentos vinculados, se listan como soporte documental al final del criterio correspondiente.",
    )
    score_by_criterion = {score.criterion_id: score for score in candidate.scores}
    file_names = {file.id: file.original_name for file in candidate.files}
    grouped: dict[str, list[Criterion]] = defaultdict(list)
    for criterion in criteria:
        grouped[criterion.category].append(criterion)

    for category, children in grouped.items():
        heading(document, category, 2)
        for criterion in children:
            score = score_by_criterion.get(criterion.id)
            body(document, f"{criterion.aspect}. Puntuación registrada: {score_text(score.score if score else None)}.")
            if score and score.rationale:
                body(document, f"Comentario de evaluación: {score.rationale}")
            if score and score.evaluator_note:
                body(document, f"Observación del evaluador: {score.evaluator_note}")
            references = [file_names.get(file_id, str(file_id)) for file_id in (score.file_ids if score else [])]
            if references:
                bullet(document, ", ".join(references))


def conclusion_text(summary: dict, template: Template) -> str:
    if summary["recommendation"] == "No califica por criterio crítico":
        return (
            "La evaluación registra al menos un criterio crítico no cumplido o no evidenciado. Conforme a la lógica "
            "definida en la plantilla, esta condición impide que el candidato califique globalmente, aun cuando pueda "
            "presentar fortalezas parciales en otras categorías."
        )
    score = summary["global_score"]
    if score >= 0.85:
        tone = "El perfil presenta una correspondencia alta con los criterios definidos."
    elif score >= 0.7:
        tone = "El perfil presenta una correspondencia favorable con los criterios definidos, con aspectos que pueden ser revisados en fases posteriores."
    elif score >= 0.55:
        tone = "El perfil requiere revisión adicional para determinar si las brechas observadas pueden ser compensadas por entrevista, examen o validación técnica."
    else:
        tone = "El perfil muestra una correspondencia limitada con los criterios definidos para la plantilla aplicada."
    return (
        f"{tone} El resultado debe analizarse junto con las evidencias documentales, las observaciones del evaluador y "
        f"los objetivos específicos de la vacante evaluada mediante la plantilla \"{template.name}\"."
    )


def build_candidate_report(candidate: Candidate, template: Template, criteria: list[Criterion]) -> BytesIO:
    document = Document()
    configure(document)
    summary = summarize_candidate(candidate, criteria)

    add_cover(document, candidate, template, summary)
    category_narrative(document, template, summary)
    weights_narrative(document, template, criteria)
    document.add_section(WD_SECTION.NEW_PAGE)
    evaluation_narrative(document, candidate, criteria)

    heading(document, "Observaciones generales")
    body(document, candidate.comments or "No se registraron observaciones generales para este candidato.")

    heading(document, "Conclusión")
    body(document, conclusion_text(summary, template))
    body(
        document,
        "Este reporte tiene carácter de soporte documental. La decisión final sobre el proceso de selección corresponde "
        "a las instancias humanas competentes, considerando el expediente completo, las entrevistas, validaciones y "
        "demás elementos institucionales aplicables.",
    )

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
